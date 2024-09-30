from django.shortcuts import render
from . models import Categories, Orders, OrderDetails, Employees
import pyodbc

from django.db.models import Q, Avg, Sum, Min, Max, Count
import csv, json, openpyxl
from io import BytesIO
from django.http import HttpResponse
from docx import Document
# from weasyprint import HTML
from django.template.loader import render_to_string
# Create your views here.
def ShowCategories(request):
    categories = Categories.objects.all()
    
    return render(request, 'dbfa/ShowCategories.html', {'categories': categories}) 

def RawSqlDemo(request):
    query = '''
    SELECT a.OrderID, a.OrderDate, b.CompanyName, c.ProductName, d.UnitPrice, d.Quantity, d.UnitPrice * d.Quantity as 'BillAmount'
    FROM orders a inner join [order details] d on a.orderid = d.orderid inner join customers b on a.customerid=b.customerid inner
    join products c on d.productid=c.productid WHERE a.orderid between 10248 and 10255
    '''
    cnxn = GetConnection()
    cursor = cnxn.cursor()
    cursor.execute(query)
    orders = cursor.fetchall()
    
    cursor.close()
    cnxn.close()
    return render(request, 'dbfa/ShowOrders.html', {'Orders': orders})

def StoredProcedureDemo(request):
    
    GrandTotal = 0
    runningTotal = 0
    runningOrderTotal = 0
    
    cnxn = GetConnection()
    cursor = cnxn.cursor()
    cursor.execute("{call USP_GetAllOrders}")
    orders = cursor.fetchall()
    
    newOrders = [] 
    previousOrderID = 0
    subtotal = 0
    
    for order in orders:
        if previousOrderID == 0:
            previousOrderID = order.OrderID
            runningTotal += order.BillAmount
            runningOrderTotal += order.BillAmount
            GrandTotal += order.BillAmount
            subtotal += order.BillAmount
            newOrders.append(pushData(order, runningTotal, runningOrderTotal))
        elif previousOrderID == order.OrderID:
            runningTotal += order.BillAmount
            runningOrderTotal += order.BillAmount
            GrandTotal += order.BillAmount
            subtotal += order.BillAmount
            newOrders.append(pushData(order, runningTotal, runningOrderTotal))
        else:
            newOrders.append(pushData(0, runningTotal, 0))
            subtotal = 0
            previousOrderID = order.OrderID
            runningOrderTotal = 0
            
            runningTotal += order.BillAmount
            runningOrderTotal += order.BillAmount
            GrandTotal += order.BillAmount
            subtotal += order.BillAmount
            newOrders.append(pushData(order, runningTotal, runningOrderTotal))
            
    newOrders.append(pushData(0, runningTotal, 0))
    cursor.close()
    cnxn.close()
        
    return render(request, 'dbfa/ShowOrders.html', {'Orders': newOrders, 'GrandTotal': GrandTotal})

def pushData(order, runningTotal, runningOrderTotal):
    if order == 0:
        return {
            'OrderID': '',
            'OrderDate': '',
            'CompanyName': '',
            'ProductName': '',
            'UnitPrice': '',
            'Quantity': '',
            'BillAmount': '',
            'RunningTotal': runningTotal,
            'RunningOrderTotal': ''
        }
    else:
        return {
            'OrderID': order.OrderID,
            'OrderDate': order.OrderDate,
            'CompanyName': order.CompanyName,
            'ProductName': order.ProductName,
            'UnitPrice': order.UnitPrice,
            'Quantity': order.Quantity,
            'BillAmount': order.BillAmount,
            'RunningTotal': runningTotal,
            'RunningOrderTotal': runningOrderTotal
        }

def SPWithOutputParametersDemo(request):
    cnxn = GetConnection()
    cursor = cnxn.cursor()
    count = 0
    cursor.execute("{call USP_GetOrdersCount(?)}", count)
    count = cursor.fetchval()
    
    cursor.execute("{call USP_GetAllOrders}")
    orders = cursor.fetchall()
    
    return render(request, 'dbfa/ShowOrders.html', {'Orders': orders, 'Count': count})

def GetConnection():
    conn = pyodbc.connect('DRIVER=ODBC Driver 17 for SQL Server;Server=.;Database=Northwind;Trusted_Connection=YES;')
    return (conn)


def FilteringQuerySetsDemo(request):
    # orders = Orders.objects.all()
    
    # orders = Orders.objects.filter(freight__gt=20)
    # orders = Orders.objects.filter(freight__gte=20)
    # orders = Orders.objects.filter(freight__lt=20)
    # orders = Orders.objects.filter(freight__lte=20)
    
    # orders = Orders.objects.filter(shipcountry__exact='Germany')
    # orders = Orders.objects.filter(shipcountry__contains='land')
    # orders = Orders.objects.filter(orderid__exact=10248)
    # orders = Orders.objects.filter(employeeid__in=[1, 3, 5])
    # orders = Orders.objects.filter(employeeid__in=[1, 3, 5]).order_by('employeeid')
    # orders = Orders.objects.filter(employeeid__in=[1, 3, 5]).order_by('-employeeid')
    
    # orders = Orders.objects.filter(shipname__startswith='A')
    # orders = Orders.objects.filter(shipname__endswith='e')
    
    # orders = Orders.objects.filter(freight__range=[10, 20])
    
    # orders = Orders.objects.filter(shipname__startswith='A') | Orders.objects.filter(freight__lt=20)
    # orders = Orders.objects.filter(shipname__startswith='A') & Orders.objects.filter(freight__lt=20)
    # orders = Orders.objects.filter(Q(shipname__startswith='A') | Q(freight__lt=20))
    # orders = Orders.objects.filter(Q(shipname__startswith='A') & Q(freight__lt=20))
    # orders = Orders.objects.filter(shipname__startswith='A', freight__lt=20)
    
    
    # orders = Orders.objects.exclude(shipname__startswith='A')
    # orders = Orders.objects.filter(~Q(shipname__startswith='A'))
    
    
    # orders = Orders.objects.all().order_by('orderid')
    # orders = Orders.objects.all().order_by('-orderid')
    # orders = Orders.objects.all().order_by('shipcountry')
    
    year = 1997
    orders = Orders.objects.filter(orderdate__year=year).order_by('-orderdate', 'employeeid')
    
    avg = Orders.objects.all().aggregate(Avg('freight'))
    max = Orders.objects.all().aggregate(Max('freight'))
    min = Orders.objects.all().aggregate(Min('freight'))
    sum = Orders.objects.all().aggregate(Sum('freight'))
    count = Orders.objects.all().aggregate(Count('freight'))
    
    my_dict = {
        'Orders': orders, 
        'avg': avg['freight__avg'],
        'count': count['freight__count'],
        'sum': sum['freight__sum'],
        'min': min['freight__min'],
        'max': max['freight__max']
    }
    
    
    # print(type(orders))
    # print(str(orders.query))
    
    return render(request, 'dbfa/FilteringDemo.html', {'Orders': my_dict})


def TwoLevelAccordionDemo(request):
    orders = Orders.objects.filter(orderid__range=[10248, 102555]).order_by('orderid')
    order_ids = [order.orderid for order in orders]
    order_details_list = OrderDetails.objects.filter(orderid__in=order_ids).order_by('orderid') 
    
    return render(request, 'dbfa/OrdersWithAccordion.html', {
        'orders' : orders,
        'order_details': order_details_list
    })
    
def MultiLevelAccordionDemo(request):
    employees_list = Employees.objects.order_by('employeeid')
    order_ids = Orders.objects.filter(employeeid__in=employees_list).values_list('orderid', flat=True).distinct()
    orders = Orders.objects.filter(orderid__in=order_ids,orderid__range=[10248,10270]).order_by('orderid')
    #orders=Orders.objects.all()
    order_ids = [order.orderid for order in orders]
    print(order_ids)
    order_details_list = OrderDetails.objects.filter(orderid__in=order_ids).order_by('orderid')
    return render(request, 'dbfa/MultiLevelAccordion.html', {  
       'employees': employees_list,    
       'orders': orders,
       'order_details': order_details_list,        
    })  




def ShowOrdersUsingCTT(request):
    return render(request, 'dbfa/ShowOrdersUsingCTT.html')




from django.views.decorators.cache import cache_page
from django.core.cache import cache


# @cache_page(60 * 15)  # Cache for 15 minutes
def CachingDemo(request):
    if cache.get('cache_Employees') == None:
        employees_list = Employees.objects.order_by('employeeid')
        
        cache.set('cache_Employees', employees_list, 3600)
        
        order_ids = Orders.objects.filter(employeeid__in=employees_list).values_list('orderid', flat=True).distinct()
        orders = Orders.objects.filter(orderid__in=order_ids,orderid__range=[10248,10270]).order_by('orderid')
        
        cache.set('cache_Orders', orders, 3600)
        #orders=Orders.objects.all()
        order_ids = [order.orderid for order in orders]
        print(order_ids)
        order_details_list = OrderDetails.objects.filter(orderid__in=order_ids).order_by('orderid')
        cache.set('cache_OrderDetails', order_details_list, 3600)
    else:
        employees_list = cache.get('cache_Employees')
        orders = cache.get('cache_Orders')
        order_details_list = cache.get('cache_OrderDetails')
        
    return render(request, 'dbfa/MultiLevelAccordion.html', {  
       'employees': employees_list,    
       'orders': orders,
       'order_details': order_details_list,        
    })  
    
    
    
def ExportToCSV(request):
    categories = Categories.objects.all()
    file_name = f'Category_data.csv'
    
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename={file_name}'
    
    writer = csv.writer(response)
    
    writer.writerow(['Category ID', 'Category Name', 'Description'])
    for category in categories:
        writer.writerow([category.categoryid, category.categoryname, category.description])
    
    return response

def ExportToJSON(request):
    categories = Categories.objects.all()
    file_name = f'Category_data.json'
    
    response = HttpResponse(content_type='application/json')
    response['Content-Disposition'] = f'attachment; filename={file_name}' 
    
    data = [{'categoryid': category.categoryid, 'categoryname': category.categoryname, 'description': category.description} for category in categories]
    
    json.dump(data, response)
    
    
    return response

def ExportToXLS(request):
    categories = Categories.objects.all()
    file_name = f'Category_data.xlsx'
    
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    
    headers = ['CategoryID', 'CategoryName', 'Description']
    worksheet.append(headers)
    
    for category in categories:
        worksheet.append([category.categoryid, category.categoryname, category.description])
    
    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename={file_name}'
    
    return response


def ExportToWord(request):
    categories=Categories.objects.all()  
    file_name=f"Category_data.docx"
     # Generate a Word document
    document = Document()
    # Add a table with headers
    table = document.add_table(rows=1, cols=3)
    table.style = 'TableGrid'
    header_row = table.rows[0].cells
    header_row[0].text = 'Category ID'
    header_row[1].text = 'Category Name'
    header_row[2].text = 'Description'
    
    for cell in header_row:
         cell.paragraphs[0].runs[0].font.bold = True
    # Add data to the table
    for category in categories:
         row = table.add_row().cells
         row[0].text = str(category.categoryid)
         row[1].text = category.categoryname
         row[2].text = category.description
    
      # Save the Word document to a BytesIO buffer
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    # Return the Word file as the HTTP response
    response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{file_name}"'
    return response

# def ExportToPDF(request):     
#     categories=Categories.objects.all()  
#     file_name=f"Category_data.pdf"
#     response = HttpResponse(content_type='application/pdf')
#     response['Content-Disposition'] = f'attachment; filename="{file_name}"'
#     html_string = render_to_string('dbfa/ShowCategoriesPDF.html', {'Categories': categories})
#     html = HTML(string=html_string)
#     html.write_pdf(response)
#     return response